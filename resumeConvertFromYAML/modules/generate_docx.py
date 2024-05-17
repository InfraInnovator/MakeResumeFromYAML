from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
import docx.shared
from docx.shared import Inches, Pt, RGBColor


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_format.space_after = Pt(0)
    p_format.space_before = Pt(0)

    # Create a new border element
    p_borders = OxmlElement('w:pBdr')

    # Create top border element
    top_bdr = OxmlElement('w:top')
    top_bdr.set(qn('w:val'), 'single')
    top_bdr.set(qn('w:sz'), '12')  # This determines the thickness of the line
    top_bdr.set(qn('w:space'), '1')
    top_bdr.set(qn('w:color'), '000000')  # Black color

    # Add top border to paragraph borders
    p_borders.append(top_bdr)

    # Add borders to paragraph
    p._element.get_or_add_pPr().append(p_borders)

    # Add a non-breaking space to the paragraph to ensure the border is visible
    run = p.add_run('\u00A0')
    run.font.color.rgb = RGBColor(255, 255, 255)  # Set font color to white to hide the non-breaking space

def add_horizontal_line_to_header(header, width):
    table = header.add_table(rows=1, cols=1, width=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = width  # Set the table width
    cell.height = docx.shared.Pt(1)  # Adjust the height for line thickness
    # cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>'))

    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    top_tcBorder = OxmlElement('w:top')
    top_tcBorder.set(qn('w:val'), 'single')
    top_tcBorder.set(qn('w:sz'), '4')
    top_tcBorder.set(qn('w:color'), '000000')
    tcBorders.append(top_tcBorder)
    tcPr.append(tcBorders)

    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_break(doc):
    doc.add_page_break()

def add_hyperlink(paragraph, text, url):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param text: The text displayed for the url.
    :param url: The url to add.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._element.append (hyperlink)

    # Add the required formatting
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def generate_docx(data, output_file):
    # Create a new Document
    doc = Document()

    # Create a new style for the header name
    header_name_style_name = 'HeaderNameStyle'
    header_name_style = doc.styles.add_style(header_name_style_name, WD_STYLE_TYPE.PARAGRAPH)
    header_name_style.base_style = doc.styles['Heading 1']
    header_name_style.font.name = 'Calibri'  # or whatever font the Title style uses
    header_name_style.font.size = Pt(26)  # adjust as needed
    header_name_style.font.bold = False
    header_name_style.font.color.rgb = RGBColor(19, 41, 75)
    header_name_style.font.underline = None
    header_name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_name_style.paragraph_format.space_before = Pt(5)
    header_name_style.paragraph_format.space_after = Pt(0)
    header_name_style.paragraph_format.line_spacing = 1.0

    # Header Styles - Subheader
    header_subheader_style = doc.styles['Subtitle']
    header_subheader_style.font.size = Pt(14)
    header_subheader_style.font.italic = False
    header_subheader_style.font.bold = True
    header_subheader_style.font.color.rgb = RGBColor(19, 41, 75)
    header_subheader_style.paragraph_format.space_before = Pt(0)
    header_subheader_style.paragraph_format.space_after = Pt(5)
    header_subheader_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_subheader_style.paragraph_format.line_spacing = 1.0

    # Header Styles - Contact
    contact_style_name = 'ContactStyle'
    header_contact_style = doc.styles.add_style(contact_style_name, WD_STYLE_TYPE.PARAGRAPH)
    header_contact_style.base_style = doc.styles['Normal']
    header_contact_style.font.size = Pt(11)
    header_contact_style.paragraph_format.space_before = Pt(0)
    header_contact_style.paragraph_format.space_after = Pt(4)
    header_contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section_heading_style = doc.styles['Heading 1']
    section_heading_style.font.size = Pt(13)
    section_heading_style.font.bold = True
    section_heading_style.font.italic = False
    section_heading_style.paragraph_format.space_before = Pt(0)
    section_heading_style.paragraph_format.space_after = Pt(0)
    section_heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    company_name_style = doc.styles['Heading 3']
    company_name_style.font.size = Pt(12)
    company_name_style.font.bold = False
    company_name_style.font.italic = False
    company_name_style.paragraph_format.space_before = Pt(0)
    company_name_style.paragraph_format.space_after = Pt(0)
    company_name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    position_dates_style = doc.styles['Heading 3']
    position_dates_style.font.size = Pt(12)
    position_dates_style.font.bold = False
    position_dates_style.font.italic = False
    position_dates_style.paragraph_format.space_before = Pt(0)
    position_dates_style.paragraph_format.space_after = Pt(0)
    position_dates_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    position_style = doc.styles['Heading 2']
    position_style.font.size = Pt(11)
    position_style.font.bold = True
    position_style.font.italic = False
    position_style.paragraph_format.space_before = Pt(0)
    position_style.paragraph_format.space_after = Pt(0)
    position_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    position_location_style = doc.styles['Heading 2']
    position_location_style.font.size = Pt(11)
    position_location_style.font.bold = False
    position_location_style.font.italic = False
    position_location_style.paragraph_format.space_before = Pt(0)
    position_location_style.paragraph_format.space_after = Pt(0)
    position_location_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Position Description
    position_description_style = doc.styles.add_style('PositionDescription', WD_STYLE_TYPE.PARAGRAPH)
    position_description_style.base_style = doc.styles['Normal']
    position_description_style.font.size = Pt(11)
    position_description_style.font.bold = False
    position_description_style.font.italic = True
    position_description_style.paragraph_format.space_before = Pt(4)
    position_description_style.paragraph_format.space_after = Pt(4)
    position_description_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    list_bullet_style = doc.styles.add_style('NewBulletStyle', WD_STYLE_TYPE.PARAGRAPH)
    list_bullet_style.base_style = doc.styles['List Bullet']
    list_bullet_style.font.size = Pt(11)
    list_bullet_style.font.bold = False
    list_bullet_style.font.italic = False
    list_bullet_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    list_bullet_style.paragraph_format.left_indent = Inches(0.5)

    skills_style = doc.styles['Normal']
    skills_style.font.size = Pt(11)
    skills_style.font.bold = False
    skills_style.font.italic = False
    skills_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Access the header of the document
    section = doc.sections[0]

    # Remove the top margin
    section.top_margin = Pt(0)
    section.header_distance = Pt(0)

    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    header = section.header

    # Ensure the header has enough paragraphs
    while len(header.paragraphs) < 3:
        header.add_paragraph()

    # Change data['header']['name'] to have upper case first letters for each word
    #data['header']['name'] = data['header']['name'].title()

    # Add content to the header
    header.paragraphs[0].add_run(data['header']['name'])
    header.paragraphs[0].style = header_name_style

    if data['header'].get('subheader'):
        header.paragraphs[1].add_run(data['header']['subheader'])
        header.paragraphs[1].style = header_subheader_style

    # Add contact info to the header
    contact_paragraph = header.add_paragraph()
    add_hyperlink(contact_paragraph, data['header']['phone'], "tel:" + data['header']['phone'])
    contact_paragraph.add_run(" | ")
    add_hyperlink(contact_paragraph, data['header']['email'], "mailto:" + data['header']['email'])
    contact_paragraph.add_run(" | ")
    add_hyperlink(contact_paragraph, data['header']['linkedin'], data['header']['linkedin'])
    contact_paragraph.style = header_contact_style

    available_width = section.page_width - section.left_margin - section.right_margin - docx.shared.Inches(.5)  # Subtracting 2 inches (1 inch from each side)

    # Add the horizontal line to the header
    add_horizontal_line_to_header(header, available_width)

    for paragraph in header.paragraphs:
        if not paragraph.text.strip():
            p = paragraph._element
            p.getparent().remove(p)

    # Professional Experience
    doc.add_paragraph("PROFESSIONAL EXPERIENCE", style=section_heading_style)
    add_horizontal_line(doc)

    for job in data['professional_experience']:
        if job.get('docx_page_break_before'):
            add_page_break(doc)

        # Create a table with invisible borders to hold company name, location, position, and dates
        table = doc.add_table(rows=2, cols=2)

        # Set table width to 100% of available space
        section = doc.sections[0]  # Assuming you're working with the first section
        available_width = section.page_width - section.left_margin - section.right_margin
        table.width = available_width
        for cell in table.columns[0].cells:
            cell.width = available_width / 2
        for cell in table.columns[1].cells:
            cell.width = available_width / 2

        # Remove table borders
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>'))

                # Set cell margins
                cell.margin_top = Inches(0)
                cell.margin_bottom = Inches(0)
                cell.margin_left = Inches(0)
                cell.margin_right = Inches(0)

        # Add company name to the left cell and location to the right cell in the first row
        company_cell = table.cell(0, 0)
        company_cell.text = job['company']
        company_cell.paragraphs[0].style = company_name_style

        dates_cell = table.cell(0, 1)
        dates_cell.text = job['dates']
        dates_cell.paragraphs[0].style = position_dates_style
        dates_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add position to the left cell and dates to the right cell in the second row
        position_cell = table.cell(1, 0)
        position_cell.text = job['position']
        position_cell.paragraphs[0].style = position_style

        location_cell = table.cell(1, 1)
        location_cell.text = job['location']
        location_cell.paragraphs[0].style = position_location_style
        location_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Job Description
        # Verifying if the job description is not empty
        if job['description']:
            doc.add_paragraph(job['description'], style='PositionDescription')

        # Duties List
        for duty in job['duties']:
            doc.add_paragraph(duty, style=list_bullet_style)

    # Additional Sections - Define a dictionary of handlers for each section
    section_handlers = {
        'skills': lambda section_data: (
            [add_page_break(doc)] if section_data.get('docx_page_break_before') else []
        ) + [
            doc.add_paragraph("SKILLS & INTERESTS", style=section_heading_style),
            add_horizontal_line(doc),
        ] + (
            [
                (lambda p: (
                    (lambda r: setattr(r, 'bold', True))(p.add_run('Skills: ')),
                    p.add_run(', '.join(section_data['skills']))
                ))(doc.add_paragraph(style=skills_style))
            ] if 'skills' in section_data and section_data['skills'] else []
        ) + (
            [
                (lambda p: (
                    (lambda r: setattr(r, 'bold', True))(p.add_run('Interests: ')),
                    p.add_run(', '.join(section_data['interests']))
                ))(doc.add_paragraph(style=skills_style))
            ] if 'interests' in section_data and section_data['interests'] else []
        ),

        'education': lambda section_data: (
            [add_page_break(doc)] if section_data.get('docx_page_break_before') else []
        ) + [
            doc.add_paragraph("EDUCATION", style=section_heading_style),
            add_horizontal_line(doc),
            *[doc.add_paragraph(school, style=list_bullet_style) for school in section_data['items']]
        ],
        'certifications': lambda section_data: (
            [add_page_break(doc)] if section_data.get('docx_page_break_before') else []
        ) + [
            doc.add_paragraph("CERTIFICATIONS AND AWARDS", style=section_heading_style),
            add_horizontal_line(doc),
            *[doc.add_paragraph(cert, style=list_bullet_style) for cert in section_data['certs']]
        ],
        'memberships': lambda section_data: (
            [add_page_break(doc)] if section_data.get('docx_page_break_before') else []
        ) + [
            doc.add_paragraph("PROFESSIONAL MEMBERSHIPS", style=section_heading_style),
            add_horizontal_line(doc),
            *[doc.add_paragraph(membership, style=list_bullet_style) for membership in section_data['memberships']]
        ],
        'summary_of_qualifications': lambda section_data: (
            [add_page_break(doc)] if section_data.get('docx_page_break_before') else []
        ) + [
            doc.add_paragraph("SUMMARY OF QUALIFICATIONS", style=section_heading_style),
            add_horizontal_line(doc),
            *[doc.add_paragraph(qualification, style=list_bullet_style) for qualification in section_data['qualifications']]
        ]
    }

    # Iterate over the keys in the data and generate the corresponding sections
    for key in data:
        if key in section_handlers:
            section_handlers[key](data[key])

    # Save the document
    doc.save(output_file)
